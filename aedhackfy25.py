import os
from openai import AzureOpenAI
from dotenv import dotenv_values
import time
from datetime import timedelta
import json
import streamlit as st
from PIL import Image
import base64
import requests
import io
from io import BytesIO
import autogen
from typing import Optional
from typing_extensions import Annotated
from streamlit import session_state as state
import azure.cognitiveservices.speech as speechsdk
from audiorecorder import audiorecorder
import pyaudio
import wave
import tempfile
import PyPDF2
from docx import Document
from streamlit_quill import st_quill
import fitz  # PyMuPDF

config = dotenv_values("env.env")

client = AzureOpenAI(
  azure_endpoint = config["AZURE_OPENAI_ENDPOINT_VISION_4o_LATEST"], 
  api_key=config["AZURE_OPENAI_KEY_VISION_4o_LATEST"],  
  api_version="2024-05-01-preview"
)

#model_name = "gpt-4-turbo"
#model_name = "gpt-35-turbo-16k"
#model_name = "gpt-4o-g"
model_name = "gpt-4o-2"

search_endpoint = config["AZURE_AI_SEARCH_ENDPOINT"]
search_key = config["AZURE_AI_SEARCH_KEY"]
#search_index=config["AZURE_AI_SEARCH_INDEX1"]
SPEECH_KEY = config['SPEECH_KEY']
SPEECH_REGION = config['SPEECH_REGION']
SPEECH_ENDPOINT = config['SPEECH_ENDPOINT']
search_index="cogsrch-index-rfp-vector"

citationtxt = ""



def extractproductinfo(user_input1, selected_optionmodel1):
    returntxt = ""

    message_text = [
    {"role":"system", "content":"""You are Tiles expert AI Agent. Be politely, and provide positive tone answers.
     extract recommendation provided and also specificiations to add to the cart.
     If not sure, ask the user to provide more information."""}, 
    {"role": "user", "content": f"""{user_input1}. Respond only recomemndation and details to add to the cart as JSON array."""}]

    response = client.chat.completions.create(
        model= selected_optionmodel1, #"gpt-4-turbo", # model = "deployment_name".
        messages=message_text,
        temperature=0.0,
        top_p=0.0,
        seed=105,
    )


    returntxt = response.choices[0].message.content
    return returntxt

def extracttextfrompdf(pdf_bytes):
    returntxt = ""

    if pdf_bytes:
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        num_pages = len(reader.pages)
        st.write(f"Number of pages in the PDF: {num_pages}")
        # Extract and display text from the first page
        if num_pages > 0:
            page = reader.pages[0]  # Get the first page
            text = page.extract_text()  # Extract text from the page
            returntxt = text

    return returntxt

def extractrfpinformation(user_input1, selected_optionmodel1, pdf_bytes):
    returntxt = ""

    rfttext = ""

    if pdf_bytes:
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        num_pages = len(reader.pages)
        st.write(f"Number of pages in the PDF: {num_pages}")
        # Extract and display text from the first page
        if num_pages > 0:
            for page_num in range(num_pages):
                page = reader.pages[page_num]  # Get each page
                text = page.extract_text()  # Extract text from the page
                rfttext += f"### Page {page_num + 1}\n{text}\n\n"  # Accumulate text from each page

    # print('RFP Text:', rfttext)

    message_text = [
    {"role":"system", "content":f"""You are RFP AI agent. Be politely, and provide positive tone answers.
     Based on the question do a detail analysis on RFP information and provide the best answers.
     Here is the RFT text tha was provided:
     {rfttext}
     please provide information based on rfp provided.
     Only provide answers from the content of the RFP.
     If not sure, ask the user to provide more information."""}, 
    {"role": "user", "content": f"""{user_input1}. Respond Only the answers with details on why the decision was made."""}]

    response = client.chat.completions.create(
        model= selected_optionmodel1, #"gpt-4-turbo", # model = "deployment_name".
        messages=message_text,
        temperature=0.0,
        top_p=0.0,
        seed=105,
    )

    returntxt = response.choices[0].message.content
    return returntxt

def extractrfpresults(user_input1, selected_optionmodel1, pdf_bytes, selected_optionsearch):
    returntxt = ""

    rfttext = ""

    if pdf_bytes:
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        num_pages = len(reader.pages)
        st.write(f"Number of pages in the PDF: {num_pages}")
        # Extract and display text from the first page
        if num_pages > 0:
            for page_num in range(num_pages):
                page = reader.pages[page_num]  # Get each page
                text = page.extract_text()  # Extract text from the page
                rfttext += f"### Page {page_num + 1}\n{text}\n\n"  # Accumulate text from each page

    # print('RFP Text:', rfttext)

    dstext = processpdfwithpromptgranite(user_input1, selected_optionmodel1, selected_optionsearch)

    message_text = [
    {"role":"system", "content":f"""You are RFP AI agent. Be politely, and provide positive tone answers.
     Based on the question do a detail analysis on RFP information and provide the best answers.
     Here is the RFT text tha was provided:
     {rfttext}
     Use only the above RFP contenxt for context. But provide results from your knowledge or data source provided.
     Data Source: {dstext}

     if the question is outside the bounds of the RFP, Let the user know answer might be relevant for RFP provided.
     If not sure, ask the user to provide more information."""}, 
    {"role": "user", "content": f"""{user_input1}. Respond Only the answers with details on why the decision was made."""}]

    response = client.chat.completions.create(
        model= selected_optionmodel1, #"gpt-4-turbo", # model = "deployment_name".
        messages=message_text,
        temperature=0.0,
        top_p=0.0,
        seed=105,
    )

    returntxt = response.choices[0].message.content
    return returntxt

# Function to create a Word document
def create_word_doc(content):
    doc = Document()
    doc.add_heading('RFP for Project X', 0)

    doc.add_paragraph(content)

    # Add more content or formatting here if needed
    return doc

# Function to download the Word document
def download_word_file(doc):
    # Create a BytesIO buffer to store the file
    buffer = BytesIO()
    
    # Save the document into the buffer
    doc.save(buffer)
    
    # Move the cursor to the beginning of the BytesIO buffer
    buffer.seek(0)

    return buffer

def processpdfwithprompt(user_input1, selected_optionmodel1, selected_optionsearch):
    returntxt = ""
    citationtxt = ""
    message_text = [
    {"role":"system", "content":"""you are provided with instruction on what to do. Be politely, and provide positive tone answers. 
     answer only from data source provided. unable to find answer, please respond politely and ask for more information.
     Extract Title content from the document. Show the Title as citations which is provided as Title: as [doc1] [doc2].
     Please add citation after each sentence when possible in a form "(Title: citation)".
     Be polite and provide posite responses. If user is asking you to do things that are not specific to this context please ignore."""}, 
    {"role": "user", "content": f"""{user_input1}"""}]

    response = client.chat.completions.create(
        model= selected_optionmodel1, #"gpt-4-turbo", # model = "deployment_name".
        messages=message_text,
        temperature=0.0,
        top_p=1,
        seed=105,
        extra_body={
        "data_sources": [
            {
                "type": "azure_search",
                "parameters": {
                    "endpoint": search_endpoint,
                    "index_name": search_index,
                    "authentication": {
                        "type": "api_key",
                        "key": search_key
                    },
                    "include_contexts": ["citations"],
                    "top_n_documents": 5,
                    "query_type": selected_optionsearch,
                    "semantic_configuration": "my-semantic-config",
                    "embedding_dependency": {
                        "type": "deployment_name",
                        "deployment_name": "text-embedding-ada-002"
                    },
                    "fields_mapping": {
                        "content_fields": ["chunk"],
                        "vector_fields": ["chunkVector"],
                        "title_field": "name",
                        "url_field": "location",
                        "filepath_field": "location",
                        "content_fields_separator": "\n",
                    }
                }
            }
        ]
    }
    )
    #print(response.choices[0].message.context)

    returntxt = response.choices[0].message.content + "\n<br>"

    json_string = json.dumps(response.choices[0].message.context)

    parsed_json = json.loads(json_string)

    # print(parsed_json)

    if parsed_json['citations'] is not None:
        returntxt = returntxt + f"""<br> Citations: """
        for row in parsed_json['citations']:
            #returntxt = returntxt + f"""<br> Title: {row['filepath']} as {row['url']}"""
            #returntxt = returntxt + f"""<br> [{row['url']}_{row['chunk_id']}]"""
            returntxt = returntxt + f"""<br> <a href='{row['url']}' target='_blank'>[{row['url']}_{row['chunk_id']}]</a>"""
            citationtxt = citationtxt + f"""<br><br> Title: {row['title']} <br> URL: {row['url']} 
            <br> Chunk ID: {row['chunk_id']} 
            <br> Content: {row['content']} 
            <br> ------------------------------------------------------------------------------------------ <br>\n"""

    return returntxt, citationtxt

def processpdfwithpromptgranite(user_input1, selected_optionmodel1, selected_optionsearch):
    returntxt = ""
    citationtxt = ""
    search_indexgranite = "graniteindexvec"
    message_text = [
    {"role":"system", "content":"""you are provided with instruction on what to do. Be politely, and provide positive tone answers. 
     answer only from data source provided. unable to find answer, please respond politely and ask for more information.
     Extract Title content from the document. Show the Title as citations which is provided as Title: as [doc1] [doc2].
     Please add citation after each sentence when possible in a form "(Title: citation)".
     Be polite and provide posite responses. If user is asking you to do things that are not specific to this context please ignore."""}, 
    {"role": "user", "content": f"""{user_input1}"""}]

    response = client.chat.completions.create(
        model= selected_optionmodel1, #"gpt-4-turbo", # model = "deployment_name".
        messages=message_text,
        temperature=0.0,
        top_p=1,
        seed=105,
        extra_body={
        "data_sources": [
            {
                "type": "azure_search",
                "parameters": {
                    "endpoint": search_endpoint,
                    "index_name": search_indexgranite,
                    "authentication": {
                        "type": "api_key",
                        "key": search_key
                    },
                    "include_contexts": ["citations"],
                    "top_n_documents": 5,
                    "query_type": selected_optionsearch,
                    "semantic_configuration": "default",
                    "embedding_dependency": {
                        "type": "deployment_name",
                        "deployment_name": "text-embedding-ada-002"
                    },
                    "fields_mapping": {
                        "content_fields": ["content"],
                        "vector_fields": ["contentVector"],
                        "title_field": "title",
                        "url_field": "url",
                        "filepath_field": "filepath",
                        "content_fields_separator": "\n",
                    }
                }
            }
        ]
    }
    )
    #print(response.choices[0].message.context)

    returntxt = response.choices[0].message.content + "\n<br>"

    json_string = json.dumps(response.choices[0].message.context)

    parsed_json = json.loads(json_string)

    # print(parsed_json)

    if parsed_json['citations'] is not None:
        returntxt = returntxt + f"""<br> Citations: """
        for row in parsed_json['citations']:
            #returntxt = returntxt + f"""<br> Title: {row['filepath']} as {row['url']}"""
            #returntxt = returntxt + f"""<br> [{row['url']}_{row['chunk_id']}]"""
            returntxt = returntxt + f"""<br> <a href='{row['url']}' target='_blank'>[{row['url']}_{row['chunk_id']}]</a>"""
            citationtxt = citationtxt + f"""<br><br> Title: {row['title']} <br> URL: {row['url']} 
            <br> Chunk ID: {row['chunk_id']} 
            <br> Content: {row['content']} 
            <br> ------------------------------------------------------------------------------------------ <br>\n"""

    return returntxt, citationtxt

# Function to convert file to base64 encoding
def file_to_base64(file):
    return base64.b64encode(file.read()).decode()

def getrfptopictorespond(user_input1, selected_optionmodel1, pdf_bytes):
    returntxt = ""

    rfttext = ""

    if pdf_bytes:
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        num_pages = len(reader.pages)
        st.write(f"Number of pages in the PDF: {num_pages}")
        # Extract and display text from the first page
        if num_pages > 0:
            for page_num in range(num_pages):
                page = reader.pages[page_num]  # Get each page
                text = page.extract_text()  # Extract text from the page
                rfttext += f"### Page {page_num + 1}\n{text}\n\n"  # Accumulate text from each page

    message_text = [
    {"role":"system", "content":f"""You are RFP and proposal expert AI Agent. Be politely, and provide positive tone answers.
     extract the topics to respond back with details as bullet point only.
     Only respond with high level topics and avoid details.
     Here is the RFP text: {rfttext}
     If not sure, ask the user to provide more information."""}, 
    {"role": "user", "content": f"""{user_input1}. Extract the topics to respond back high level bullet point only."""}]

    response = client.chat.completions.create(
        model= selected_optionmodel1, #"gpt-4-turbo", # model = "deployment_name".
        messages=message_text,
        temperature=0.0,
        top_p=0.0,
        seed=105,
    )


    returntxt = response.choices[0].message.content
    return returntxt
# Initialize session state for the editor content
if "quill_rfpcontent" not in st.session_state:
    st.session_state.quill_rfpcontent = ""

# Initialize session state for the editor content
if "quill_rfpresponse" not in st.session_state:
    st.session_state.quill_rfpresponse = ""

# Define a function to update the content programmatically
def update_quill_rfpcontent(new_content):
  
    # Append new content to the existing content changed
    st.session_state.quill_rfpcontent = new_content

# Define a function to update the content programmatically
def update_quill_rfpresponse(new_content):
    st.session_state.quill_rfpresponse += new_content

# Define the Linked List Node class
class Node:
    def __init__(self, topic_name, content):
        self.topic_name = topic_name
        self.content = content
        self.next = None

# Define the Linked List class
class LinkedList:
    def __init__(self):
        self.head = None

    # Method to add new node to the list
    def add(self, topic_name, content):
        # Check if the topic already exists
        if self.contains(topic_name):
            # Update the content of the existing topic
            self.update(topic_name, content)
        else:
            # Add a new node if the topic does not exist
            new_node = Node(topic_name, content)
            if not self.head:
                self.head = new_node
            else:
                current = self.head
                while current.next:
                    current = current.next
                current.next = new_node
   
    # Method to display the list
    """
    def display(self):
        current = self.head
        while current:
            st.write(f"**Topic:** {current.topic_name}")
            st.write(f"**Content:** {current.content}")
            st.write("---")
            current = current.next
   """
    def display(self):
        current = self.head
        content = ""
        while current:
            content += f"**Topic:** {current.topic_name}\n**Content:** {current.content}\n\n"
            current = current.next
        return content
    def contains(self, topic_name):
        current = self.head
        while current:
            if current.topic_name == topic_name:
                return True
            current = current.next
        return False

    def update(self, topic_name, content):
        current = self.head
        while current:
            if current.topic_name == topic_name:
                current.content = content
                return
            current = current.next

def display_pdf_as_iframe(file_path):
    with open(file_path, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
        pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="700" height="900" type="application/pdf"></iframe>'
        st.markdown(pdf_display, unsafe_allow_html=True)

# Function to convert PDF pages into images
def pdf_to_images(pdf_path, zoom=2.0):
    # Open the PDF file
    pdf_document = fitz.open(pdf_path)
    images = []

    # Iterate through all the pages
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)  # Load page
        pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))  # Render the page as an image
        image = Image.open(io.BytesIO(pix.tobytes("png")))  # Convert to PIL Image
        images.append(image)

    return images

def encode_image(image_path):
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode('utf-8')
    
def processimage(base64_image, imgprompt):
    response = client.chat.completions.create(
    model="gpt-4o",
    messages=[
        {
        "role": "user",
        "content": [
            {"type": "text", "text": f"{imgprompt}"},
            {
            "type": "image_url",
            "image_url": {
                "url" : f"data:image/jpeg;base64,{base64_image}",
            },
            },
        ],
        }
    ],
    max_tokens=2000,
    temperature=0,
    top_p=1,
    )

    #print(response.choices[0].message.content)
    return response.choices[0].message.content

def process_image(uploaded_file, selected_optionmodel, user_input):
    returntxt = ""

    if uploaded_file is not None:
        #image = Image.open(os.path.join(os.getcwd(),"temp.jpeg"))
        img_path = os.path.join(os.getcwd(), uploaded_file)
        # Open the image using PIL
        #image_bytes = uploaded_file.read()
        #image = Image.open(io.BytesIO(image_bytes))

        base64_image = encode_image(img_path)
        #base64_image = base64.b64encode(uploaded_file).decode('utf-8') #uploaded_image.convert('L')
        imgprompt = f"""You are a Constructon drawing AutoCad Expert Agent. Analyze the image and find details for questions asked.
        Only answer from the data source provided.
        Image has information about drawingsprovided.
        can you extract details of this drawings.

        Question:
        {user_input} 
        """

        # Get the response from the model
        result = processimage(base64_image, imgprompt)

        #returntxt += f"Image uploaded: {uploaded_file.name}\n"
        returntxt = result

    return returntxt

def compare_rfq_drawings(uploaded_file, selected_optionmodel, user_input, pdf_file):
    returntxt = ""
    pdftext = ""

    if pdf_file is not None:
        pdftext = extracttextfrompdf(pdf_file)

    if uploaded_file is not None:
        #image = Image.open(os.path.join(os.getcwd(),"temp.jpeg"))
        img_path = os.path.join(os.getcwd(), uploaded_file)
        # Open the image using PIL
        #image_bytes = uploaded_file.read()
        #image = Image.open(io.BytesIO(image_bytes))

        base64_image = encode_image(img_path)
        #base64_image = base64.b64encode(uploaded_file).decode('utf-8') #uploaded_image.convert('L')
        imgprompt = f"""You are a Constructon drawing AutoCad Expert Agent. Analyze the image and find details for questions asked.
        Only answer from the data source provided.
        Image has information about drawingsprovided.
        Compare the insights from drawings to RFQ provided. Provide details on if the drawings have everything needed for RFQ.
        Point out any missing details.
        Also provide recommendation on any improvements we can do.

        RFQ Text:
        {pdftext}

        Question:
        {user_input} 
        """

        # Get the response from the model
        result = processimage(base64_image, imgprompt)

        #returntxt += f"Image uploaded: {uploaded_file.name}\n"
        returntxt = result

    return returntxt

def aechackfy25():
    count = 0
    temp_file_path = ""
    pdf_bytes = None
    rfpcontent = {}
    rfplist = []
    #tab1, tab2, tab3, tab4 = st.tabs('RFP PDF', 'RFP Research', 'Draft', 'Create Word')
    modeloptions1 = ["gpt-4o-2", "gpt-4o-g", "gpt-4o", "gpt-4-turbo", "gpt-35-turbo"]



    # Create a dropdown menu using selectbox method
    selected_optionmodel1 = st.selectbox("Select an Model:", modeloptions1)
    count += 1

    tabs = st.tabs(["RFP PDF", "RFP Research", "RFP Draft", "Create Word", 
                    "Existing RFP Drawings", "Compare Drawings to RFQ"])


    with tabs[0]:
        st.write("Upload RFP PDF file")
        uploaded_file = st.file_uploader("Choose a PDF file", type="pdf", key="pdf_file0")
        if uploaded_file is not None:
            # Display the PDF in an iframe
            pdf_bytes = uploaded_file.read()  # Read the PDF as bytes
            st.download_button("Download PDF", pdf_bytes, file_name="uploaded_pdf.pdf")

            # Convert to base64
            base64_pdf = base64.b64encode(pdf_bytes).decode('utf-8')
            # Embedding PDF using an HTML iframe
            pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="1000" height="700" type="application/pdf"></iframe>'
            st.markdown(pdf_display, unsafe_allow_html=True)
            # Save the PDF file to the current folder
            file_path = os.path.join(os.getcwd(), uploaded_file.name)  # Save in the current directory
            with open(file_path, "wb") as f:
                f.write(uploaded_file.read())  # Write the uploaded file to disk
            
            # Display the path where the file is stored
            # st.write(f"File saved to: {file_path}")
            temp_file_path = file_path

    with tabs[1]:
        st.write("Search for information in RFP")        
        query = st.text_input("Enter your search query", "Summarize the content of the RFP")
        
        if st.button("Search"):
            # Call the extractproductinfo function
            result = extractrfpinformation(query, selected_optionmodel1, pdf_bytes)
            # st.text_input("Output", result)
            # st.markdown(result, unsafe_allow_html=True)
            #st.session_state.quill_rfpcontent = result
            update_quill_rfpcontent(result)
            # print('RFP Content:', result)
            #rfpcontent1 = st_quill(result)
            #rfpcontent1
        rfpcontent1 = st_quill(st.session_state.quill_rfpcontent, placeholder="Enter your rich text...",    key="editor1")
    with tabs[2]:
        st.write("RFP Topic List")
        col1, col2 = st.columns([1, 2])
        with col1:
            rfpquery = st.text_input("Enter your RFP query", "what are qualification or critical success to winning this RFQ")
            selected_optionsearch = st.selectbox("Select Search Type", ["simple", "semantic", "vector", "vector_simple_hybrid", "vector_semantic_hybrid"])
            topic_name = st.text_input("Enter the topic name:", "Introduction")
            # Call the extractproductinfo function
            rfttopics = getrfptopictorespond(rfpquery, selected_optionmodel1, pdf_bytes)
            st.markdown(rfttopics)
            rfpcontent = {"topic": "rftcontent", "result": rfttopics}
            st.image("pavement1.jpg", use_column_width=True)
        #rfttopics = getrfptopictorespond(query, selected_optionmodel1, pdf_bytes)
        #st.markdown(rfttopics)
        #st.write("RFP Draft")
        #rfpquery = st.text_input("Enter your RFP query", "show me insights on bridge construction")
        with col2:                 
            if st.button("rfp content"):
                # Call the extractproductinfo function
                result = extractrfpresults(rfpquery, selected_optionmodel1, pdf_bytes, selected_optionsearch)
                #result = processpdfwithprompt(rfpquery, selected_optionmodel1, selected_optionsearch)
                #st.text_input(value=result)
                #st.text_input("Output", result)
                # st.markdown(result, unsafe_allow_html=True)
                # quill_rfpresponse
                #st.session_state.quill_rfpresponse = result changed
                
                update_quill_rfpresponse(result)
               

                rfpcontent = {"topic": "rftcontent", "result": result}
            rfpresponse1 = st_quill(st.session_state.quill_rfpresponse, placeholder="Enter your rich text here...",    key="editor_rfp_response")
            update_quill_rfpresponse(rfpresponse1)
            # Add a save button
            if st.button("Save"):
               # sections = st.session_state.quill_rfpresponse.strip().split('\n\n')
                sections = st.session_state.quill_rfpresponse.strip().split('\n\n')
                # Initialize the linked list
                # Initialize session state for storing the linked list
                if "linked_list" not in st.session_state:
                    st.session_state.linked_list = LinkedList()

                # Add each section to the linked list
                for section in sections:
                    if '**' in section:
                        parts = section.split('**')
                        if len(parts) >= 3:
                            topic = parts[1].strip()
                            content = parts[2].strip()
                            st.session_state.linked_list.add(topic, content)
                    else:
                        st.session_state.linked_list.add('Introduction', section)
                st.text_area("RFP Response",  st.session_state.linked_list.display(), height=400, key="editor_linked_list")
    with tabs[3]:
        st.write("Create Word Document")
        # result = extractrfpresults(rfpquery, selected_optionmodel1, pdf_bytes, selected_optionsearch)
        if rfpresponse1:
            st.session_state.quill_rfpresponse = rfpresponse1
            resultdoc = st.session_state.quill_rfpresponse
            #st.write(result)
            rfpcontent = {"topic": "RFP Content", "result": resultdoc}

        if st.button("Generate and Download Word Document"):
            if rfpcontent:
                rfptopic = rfpcontent["topic"]
                user_input = rfpcontent["result"]
                #print('User Input:', user_input)
                # Create the Word document
                doc = create_word_doc(user_input)
                #print('Doc:', doc)
                
                # Get the document in a buffer
                word_file = download_word_file(doc)
                
                # Create a download button
                st.download_button(
                    label="Download Word Document",
                    data=word_file,
                    file_name="generated_document.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                # base64_pdf = base64.b64encode(word_file).decode('utf-8')
                # Embed the Word document using an iframe
                #st.markdown(f"""
                #    <iframe src="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{word_file}" 
                #            width="700" height="500" type="application/vnd.openxmlformats-officedocument.wordprocessingml.document">
                #    </iframe>
                #""", unsafe_allow_html=True)
        else:
            st.error("Check your rfp content.")
    with tabs[4]:
        st.write("Displaying Existing RFP Drawing")
        col1, col2 = st.columns([1, 2])

        with col1:
            st.write("Upload RFP Drawing")
            constimgfile = "constr2.jpg"
            drawingtext = process_image(constimgfile, selected_optionmodel1, "what are the details of this drawing")
            st.markdown(drawingtext, unsafe_allow_html=True)
            
        with col2:
            # uploaded_file1 = os.getcwd() + "/Preliminary_Plans.pdf"
            pdf_file = os.getcwd() + "\\Preliminary_Plans.pdf"
            print(pdf_file)
            # Convert PDF pages to images
            images = pdf_to_images(pdf_file, zoom=2.0)

            # Save the images or show them (optional)
            #for i, img in enumerate(images):
            #    img.save(f"page_{i+1}.png")  # Save each page as an image file
            #    img.show()  # To display the image

            for i, img in enumerate(images):
                st.image(img, caption=f"Page {i+1}", use_column_width=True)
    with tabs[5]:
        st.write("Compare Drawings to RFQ")
        constimgfile = "constr2.jpg"
        if uploaded_file:
            pdf_bytes2 = uploaded_file.read()  # Read the PDF as bytes
            drwaingsinsightstext = compare_rfq_drawings(constimgfile, selected_optionmodel1, "what are the details of this drawing", pdf_bytes2)
            st.markdown(drwaingsinsightstext, unsafe_allow_html=True)