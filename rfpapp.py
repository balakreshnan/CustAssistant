import os
from openai import AzureOpenAI
from dotenv import dotenv_values
import time
from datetime import timedelta
import json
import streamlit as st
from PIL import Image
import base64
import requests
import io
from io import BytesIO
import autogen
from typing import Optional
from typing_extensions import Annotated
from streamlit import session_state as state
import azure.cognitiveservices.speech as speechsdk
from audiorecorder import audiorecorder
import pyaudio
import wave
import tempfile
import PyPDF2
from docx import Document

config = dotenv_values("env.env")

client = AzureOpenAI(
  azure_endpoint = config["AZURE_OPENAI_ENDPOINT_VISION_4o_LATEST"], 
  api_key=config["AZURE_OPENAI_KEY_VISION_4o_LATEST"],  
  api_version="2024-05-01-preview"
)

#model_name = "gpt-4-turbo"
#model_name = "gpt-35-turbo-16k"
#model_name = "gpt-4o-g"
model_name = "gpt-4o-2"

search_endpoint = config["AZURE_AI_SEARCH_ENDPOINT"]
search_key = config["AZURE_AI_SEARCH_KEY"]
#search_index=config["AZURE_AI_SEARCH_INDEX1"]
SPEECH_KEY = config['SPEECH_KEY']
SPEECH_REGION = config['SPEECH_REGION']
SPEECH_ENDPOINT = config['SPEECH_ENDPOINT']
search_index="cogsrch-index-rfp-vector"

citationtxt = ""

def extractproductinfo(user_input1, selected_optionmodel1):
    returntxt = ""

    message_text = [
    {"role":"system", "content":"""You are Tiles expert AI Agent. Be politely, and provide positive tone answers.
     extract recommendation provided and also specificiations to add to the cart.
     If not sure, ask the user to provide more information."""}, 
    {"role": "user", "content": f"""{user_input1}. Respond only recomemndation and details to add to the cart as JSON array."""}]

    response = client.chat.completions.create(
        model= selected_optionmodel1, #"gpt-4-turbo", # model = "deployment_name".
        messages=message_text,
        temperature=0.0,
        top_p=0.0,
        seed=105,
    )


    returntxt = response.choices[0].message.content
    return returntxt

def extractrfpinformation(user_input1, selected_optionmodel1, pdf_bytes):
    returntxt = ""

    rfttext = ""

    if pdf_bytes:
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        num_pages = len(reader.pages)
        st.write(f"Number of pages in the PDF: {num_pages}")
        # Extract and display text from the first page
        if num_pages > 0:
            for page_num in range(num_pages):
                page = reader.pages[page_num]  # Get each page
                text = page.extract_text()  # Extract text from the page
                rfttext += f"### Page {page_num + 1}\n{text}\n\n"  # Accumulate text from each page

    # print('RFP Text:', rfttext)

    message_text = [
    {"role":"system", "content":f"""You are RFP AI agent. Be politely, and provide positive tone answers.
     Based on the question do a detail analysis on RFP information and provide the best answers.
     Here is the RFT text tha was provided:
     {rfttext}
     please provide information based on rfp provided.
     Only provide answers from the content of the RFP.
     If not sure, ask the user to provide more information."""}, 
    {"role": "user", "content": f"""{user_input1}. Respond Only the answers with details on why the decision was made."""}]

    response = client.chat.completions.create(
        model= selected_optionmodel1, #"gpt-4-turbo", # model = "deployment_name".
        messages=message_text,
        temperature=0.0,
        top_p=0.0,
        seed=105,
    )

    returntxt = response.choices[0].message.content
    return returntxt

def extractrfpresults(user_input1, selected_optionmodel1, pdf_bytes, selected_optionsearch):
    returntxt = ""

    rfttext = ""

    if pdf_bytes:
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        num_pages = len(reader.pages)
        st.write(f"Number of pages in the PDF: {num_pages}")
        # Extract and display text from the first page
        if num_pages > 0:
            for page_num in range(num_pages):
                page = reader.pages[page_num]  # Get each page
                text = page.extract_text()  # Extract text from the page
                rfttext += f"### Page {page_num + 1}\n{text}\n\n"  # Accumulate text from each page

    # print('RFP Text:', rfttext)

    dstext = processpdfwithprompt(user_input1, selected_optionmodel1, selected_optionsearch)

    message_text = [
    {"role":"system", "content":f"""You are RFP AI agent. Be politely, and provide positive tone answers.
     Based on the question do a detail analysis on RFP information and provide the best answers.
     Here is the RFT text tha was provided:
     {rfttext}
     Use only the above RFP contenxt for context. But provide results from your knowledge or data source provided.
     Data Source: {dstext}

     if the question is outside the bounds of the RFP, Let the user know answer might be relevant for RFP provided.
     If not sure, ask the user to provide more information."""}, 
    {"role": "user", "content": f"""{user_input1}. Respond Only the answers with details on why the decision was made."""}]

    response = client.chat.completions.create(
        model= selected_optionmodel1, #"gpt-4-turbo", # model = "deployment_name".
        messages=message_text,
        temperature=0.0,
        top_p=0.0,
        seed=105,
    )

    returntxt = response.choices[0].message.content
    return returntxt

# Function to create a Word document
def create_word_doc(content):
    doc = Document()
    doc.add_heading('Generated Word Document', 0)

    doc.add_paragraph(content)

    # Add more content or formatting here if needed
    return doc

# Function to download the Word document
def download_word_file(doc):
    # Create a BytesIO buffer to store the file
    buffer = BytesIO()
    
    # Save the document into the buffer
    doc.save(buffer)
    
    # Move the cursor to the beginning of the BytesIO buffer
    buffer.seek(0)

    return buffer

def processpdfwithprompt(user_input1, selected_optionmodel1, selected_optionsearch):
    returntxt = ""
    citationtxt = ""
    message_text = [
    {"role":"system", "content":"""you are provided with instruction on what to do. Be politely, and provide positive tone answers. 
     answer only from data source provided. unable to find answer, please respond politely and ask for more information.
     Extract Title content from the document. Show the Title as citations which is provided as Title: as [doc1] [doc2].
     Please add citation after each sentence when possible in a form "(Title: citation)".
     Be polite and provide posite responses. If user is asking you to do things that are not specific to this context please ignore."""}, 
    {"role": "user", "content": f"""{user_input1}"""}]

    response = client.chat.completions.create(
        model= selected_optionmodel1, #"gpt-4-turbo", # model = "deployment_name".
        messages=message_text,
        temperature=0.0,
        top_p=1,
        seed=105,
        extra_body={
        "data_sources": [
            {
                "type": "azure_search",
                "parameters": {
                    "endpoint": search_endpoint,
                    "index_name": search_index,
                    "authentication": {
                        "type": "api_key",
                        "key": search_key
                    },
                    "include_contexts": ["citations"],
                    "top_n_documents": 5,
                    "query_type": selected_optionsearch,
                    "semantic_configuration": "my-semantic-config",
                    "embedding_dependency": {
                        "type": "deployment_name",
                        "deployment_name": "text-embedding-ada-002"
                    },
                    "fields_mapping": {
                        "content_fields": ["chunk"],
                        "vector_fields": ["chunkVector"],
                        "title_field": "name",
                        "url_field": "location",
                        "filepath_field": "location",
                        "content_fields_separator": "\n",
                    }
                }
            }
        ]
    }
    )
    #print(response.choices[0].message.context)

    returntxt = response.choices[0].message.content + "\n<br>"

    json_string = json.dumps(response.choices[0].message.context)

    parsed_json = json.loads(json_string)

    # print(parsed_json)

    if parsed_json['citations'] is not None:
        returntxt = returntxt + f"""<br> Citations: """
        for row in parsed_json['citations']:
            #returntxt = returntxt + f"""<br> Title: {row['filepath']} as {row['url']}"""
            #returntxt = returntxt + f"""<br> [{row['url']}_{row['chunk_id']}]"""
            returntxt = returntxt + f"""<br> <a href='{row['url']}' target='_blank'>[{row['url']}_{row['chunk_id']}]</a>"""
            citationtxt = citationtxt + f"""<br><br> Title: {row['title']} <br> URL: {row['url']} 
            <br> Chunk ID: {row['chunk_id']} 
            <br> Content: {row['content']} 
            <br> ------------------------------------------------------------------------------------------ <br>\n"""

    return returntxt, citationtxt

def showrfpoptions():
    count = 0
    temp_file_path = ""
    pdf_bytes = None
    rfpcontent = {}
    rfplist = []
    #tab1, tab2, tab3, tab4 = st.tabs('RFP PDF', 'RFP Research', 'Draft', 'Create Word')
    modeloptions1 = ["gpt-4o-2", "gpt-4o-g", "gpt-4o", "gpt-4-turbo", "gpt-35-turbo"]

    # Create a dropdown menu using selectbox method
    selected_optionmodel1 = st.selectbox("Select an Model:", modeloptions1)
    count += 1

    tabs = st.tabs(["RFP PDF", "RFP Research", "RFP Draft", "Create Word"])

    with tabs[0]:
        st.write("Upload RFP PDF file")
        uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")
        if uploaded_file is not None:
            # Display the PDF in an iframe
            pdf_bytes = uploaded_file.read()  # Read the PDF as bytes
            st.download_button("Download PDF", pdf_bytes, file_name="uploaded_pdf.pdf")

            # Convert to base64
            base64_pdf = base64.b64encode(pdf_bytes).decode('utf-8')
            # Embedding PDF using an HTML iframe
            pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="700" height="1000" type="application/pdf"></iframe>'
            st.markdown(pdf_display, unsafe_allow_html=True)
            # Save the PDF file to the current folder
            file_path = os.path.join(os.getcwd(), uploaded_file.name)  # Save in the current directory
            with open(file_path, "wb") as f:
                f.write(uploaded_file.read())  # Write the uploaded file to disk
            
            # Display the path where the file is stored
            st.write(f"File saved to: {file_path}")
            temp_file_path = file_path

    with tabs[1]:
        st.write("Search for information in RFP")        
        query = st.text_input("Enter your search query")
        if st.button("Search"):
            # Call the extractproductinfo function
            result = extractrfpinformation(query, selected_optionmodel1, pdf_bytes)
            st.write(result)
    with tabs[2]:
        st.write("Draft RFP")
        rfpquery = st.text_input("Enter your RFP query")
        selected_optionsearch = st.selectbox("Select Search Type", ["simple", "semantic", "vector", "vector_simple_hybrid", "vector_semantic_hybrid"])
        if st.button("rfp content"):
            # Call the extractproductinfo function
            result = extractrfpresults(rfpquery, selected_optionmodel1, pdf_bytes, selected_optionsearch)
            #result = processpdfwithprompt(rfpquery, selected_optionmodel1, selected_optionsearch)
            st.write(result)
            rfpcontent = {"topic": "rftcontent", "result": result}

    with tabs[3]:
        st.write("Create Word Document")
        result = extractrfpresults(rfpquery, selected_optionmodel1, pdf_bytes, selected_optionsearch)
        #st.write(result)
        rfpcontent = {"topic": "rftcontent", "result": result}

        if st.button("Generate and Download Word Document"):
            if rfpcontent:
                rfptopic = rfpcontent["topic"]
                user_input = rfpcontent["result"]
                #print('User Input:', user_input)
                # Create the Word document
                doc = create_word_doc(user_input)
                #print('Doc:', doc)
                
                # Get the document in a buffer
                word_file = download_word_file(doc)
                
                # Create a download button
                st.download_button(
                    label="Download Word Document",
                    data=word_file,
                    file_name="generated_document.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.error("Check your rfp content.")